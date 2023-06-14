from datetime import datetime, date
from io import BytesIO
from ..models import *
from ..serializers.organization_serializer import OrganizationSerializer
from rest_framework import status
from rest_framework.response import Response
from docx import Document

def get_organization(organization_id, request):
    if request.user.organization_id != organization_id:
        return Response({'message': "User can get only his/her organization"}, status.HTTP_403_FORBIDDEN)
    try:
        instance = Organization.objects.get(pk=organization_id)
        organization_serializer = OrganizationSerializer(instance)
    except Exception as e:
        return Response({'message': str(e)}, status.HTTP_204_NO_CONTENT)
    return Response(organization_serializer.data, status.HTTP_200_OK)

def get_organization_statistics(organization_id, request):
    if request.user.organization_id != organization_id:
        return Response({'message': "User can get only his/her organization statistics"}, status.HTTP_403_FORBIDDEN)
    try:

        bracelets = Bracelet.objects.filter(organization_id=organization_id)
        total_bracelets = bracelets.count()

        today = date.today()
        customUsers = CustomUser.objects.filter(organization_id=organization_id)
        
        actual_bracelets = bracelets.filter(status=True)
        user_ids = actual_bracelets.values('customuser_id')
        actual_users = customUsers.filter(id__in=user_ids)
        visitors = actual_users.filter(is_active=True)
        current_visitors = visitors.filter(role_id=2)
        current_visitors_count = current_visitors.count()

        zones = Zone.objects.filter(organization_id=organization_id)
        total_zones = zones.count()

        services = Service.objects.filter(organization_id=organization_id)
        total_services = services.count()

        bracelet_ids = bracelets.values('id')
        interactions = Interaction.objects.filter(id__in=bracelet_ids)
        total_interactions_today = interactions.filter(interaction_datetime__date=today).count()
        
        statistics = {
            'currentVisitorsCount': current_visitors_count,
            'totalBracelets': total_bracelets,
            'totalZones': total_zones,
            'totalServices': total_services,
            'totalInteractionsToday': total_interactions_today
        }
    except Exception as e:
        return Response({'message': str(e)}, status.HTTP_204_NO_CONTENT)
    return Response(statistics, status.HTTP_200_OK)

def generate_word_report(organization_id, request):
    if request.user.organization_id != organization_id:
        return Response({'message': "User can get only his/her organization statistics"}, status.HTTP_403_FORBIDDEN)
    try:

        bracelets = Bracelet.objects.filter(organization_id=organization_id)
        total_bracelets = bracelets.count()

        today = date.today()
        customUsers = CustomUser.objects.filter(organization_id=organization_id)
        
        actual_bracelets = bracelets.filter(status=True)
        user_ids = actual_bracelets.values('customuser_id')
        actual_users = customUsers.filter(id__in=user_ids)
        visitors = actual_users.filter(is_active=True)
        current_visitors = visitors.filter(role_id=2)
        current_visitors_count = current_visitors.count()

        zones = Zone.objects.filter(organization_id=organization_id)
        total_zones = zones.count()

        services = Service.objects.filter(organization_id=organization_id)
        total_services = services.count()

        bracelet_ids = bracelets.values('id')
        interactions = Interaction.objects.filter(id__in=bracelet_ids)
        total_interactions_today = interactions.filter(interaction_datetime__date=today).count()
        organization = Organization.objects.get(pk=organization_id)
        statistics = {
            'name': organization.name,
            'address': organization.address,
            'email': organization.email,
            'phoneNumber': organization.phone_number,
            'currentVisitorsCount': current_visitors_count,
            'totalBracelets': total_bracelets,
            'totalZones': total_zones,
            'totalServices': total_services,
            'totalInteractionsToday': total_interactions_today
        }

        document = Document()

        # Додавання заголовка
        document.add_heading('Statistics Report', level=1)

        # Додавання розділу для загальної інформації
        document.add_heading('General Information', level=2)
        document.add_paragraph(f"Organization Name: {statistics.get('name')}")
        document.add_paragraph(f"Address: {statistics.get('address')}")
        document.add_paragraph(f"Email: {statistics.get('email')}")
        document.add_paragraph(f"Phone: {statistics.get('phoneNumber')}")
        document.add_paragraph(f"Current Visitors Count: {statistics.get('currentVisitorsCount')}")
        document.add_paragraph(f"Total Bracelets: {statistics.get('totalBracelets')}")
        document.add_paragraph(f"Total Zones: {statistics.get('totalZones')}")
        document.add_paragraph(f"Total Services: {statistics.get('totalServices')}")
        document.add_paragraph(f"Total Interactions Today: {statistics.get('totalInteractionsToday')}")

            # Збереження документу в пам'яті
        output = BytesIO()
        document.save(output)
        output.seek(0)

        # Create the response object
        response = Response(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="report.docx"'

        # Set the response content with the document data
        response.content = output.getvalue()

        # Return the response
        return response
        # return Response(
        #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        #     content=output.getvalue(),
        #     headers={'Content-Disposition': 'attachment; filename="report.docx"'},
        #     status=status.HTTP_200_OK
        # )
    except Exception as e:
        # Handle any exceptions and return an error response
        error_message = str(e)
        return Response({'message': str(e)}, status.HTTP_204_NO_CONTENT)


def get_organizations():
    organizations = Organization.objects.all()
    serializer = OrganizationSerializer(organizations, many = True)
    if len(serializer.data) == 0:
        return Response(serializer.data, status.HTTP_204_NO_CONTENT)
    return Response(serializer.data, status.HTTP_200_OK)

def create_organization(data):
    serializer = OrganizationSerializer(data=data)
    
    if serializer.is_valid():
        try:
            serializer.save()            
            return Response(serializer.data, status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'message': str(e)}, status.HTTP_400_BAD_REQUEST)
    return Response(serializer.errors, status.HTTP_400_BAD_REQUEST)

def update_organization(organization_id, data, request):
    if request.user.organization_id != organization_id:
        return Response({'message': "User can update only his/her organization"}, status.HTTP_403_FORBIDDEN)
    try:
        organization = Organization.objects.get(pk=organization_id)
    except Exception as e:
        return Response({'message': str(e)}, status.HTTP_404_NOT_FOUND)
    
    serializer = OrganizationSerializer(organization, data=data)
    if serializer.is_valid():
        try:
            serializer.save()
        except:
            return Response({'message': 'Something went wrong'}, status.HTTP_400_BAD_REQUEST)
        
        return Response(status=status.HTTP_204_NO_CONTENT)

    return Response(serializer.errors, status.HTTP_400_BAD_REQUEST)

def delete_organization(organization_id, request):
    if request.user.organization_id != organization_id:
        return Response({'message': "User can delete only his/her organization"}, status.HTTP_403_FORBIDDEN)
    try:
        organization = Organization.objects.get(pk=organization_id)
    except Exception as e:
        return Response({'message': str(e)}, status.HTTP_404_NOT_FOUND)
    
    try:
        organization.delete()
    except:
        return Response({'message': 'Something went wrong'}, status.HTTP_400_BAD_REQUEST)
    
    return Response(status=status.HTTP_204_NO_CONTENT)